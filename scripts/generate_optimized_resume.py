from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = Path(r"D:\AIAgent\RAG\谢敬南-本科-哈尔滨理工大学-2027-优化版.docx")


def set_run_font(run, font_name="Microsoft YaHei", size=10.5, bold=False, color=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "nil")
        borders.append(elem)
    tbl_pr.append(borders)


def add_paragraph_text(container, text, *, size=10.5, bold=False, color=None, align=None, space_after=0):
    p = container.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_bullet(container, title, text):
    p = container.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.2)
    p.paragraph_format.first_line_indent = Cm(-0.45)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.12
    bullet = p.add_run("• ")
    set_run_font(bullet, size=10.5, bold=True, color=RGBColor(54, 54, 54))
    title_run = p.add_run(title)
    set_run_font(title_run, size=10.5, bold=True, color=RGBColor(54, 54, 54))
    body_run = p.add_run(text)
    set_run_font(body_run, size=10.5, color=RGBColor(54, 54, 54))


def add_section_heading(container, title):
    p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    set_run_font(run, size=12.5, bold=True, color=RGBColor(39, 63, 130))
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "8B96C5")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    return p


def build_document():
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.1)
    section.bottom_margin = Cm(1.1)
    section.left_margin = Cm(1.25)
    section.right_margin = Cm(1.25)

    style = document.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(10.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("谢敬南")
    set_run_font(run, size=18, bold=True, color=RGBColor(32, 45, 91))

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(0)
    run = subtitle.add_run("本科 | 哈尔滨理工大学 | 机器人工程 | 预计 2027.06 毕业")
    set_run_font(run, size=10.5, color=RGBColor(85, 85, 85))

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(6)
    run = contact.add_run("电话：18651522512   |   邮箱：18651522512@163.com   |   英语：CET-6")
    set_run_font(run, size=10.5, color=RGBColor(85, 85, 85))

    add_section_heading(document, "教育背景")
    edu = document.add_paragraph()
    edu.paragraph_format.space_after = Pt(2)
    edu.paragraph_format.line_spacing = 1.12
    r1 = edu.add_run("哈尔滨理工大学  ")
    set_run_font(r1, size=11, bold=True, color=RGBColor(44, 44, 44))
    r2 = edu.add_run("机器人工程 本科")
    set_run_font(r2, size=10.5, color=RGBColor(44, 44, 44))
    r3 = edu.add_run("    2027.06 毕业")
    set_run_font(r3, size=10.5, color=RGBColor(90, 90, 90))

    add_section_heading(document, "项目经历")

    p1 = document.add_paragraph()
    p1.paragraph_format.space_after = Pt(2)
    p1_run1 = p1.add_run("多文档问答系统（RAG）")
    set_run_font(p1_run1, size=11.5, bold=True, color=RGBColor(35, 52, 105))
    p1_run2 = p1.add_run("   FastAPI / FAISS / BGE / Vue 3 / SQLAlchemy / Docker")
    set_run_font(p1_run2, size=9.8, color=RGBColor(96, 96, 96))

    add_bullet(document, "", "独立完成面向本地知识库场景的文档问答系统，支持 PDF、Word、Markdown、TXT 文档上传、解析、索引与问答。")
    add_bullet(document, "", "搭建 BGE 向量检索、FAISS 召回与 Reranker 重排链路，支持按文档范围过滤并返回引用片段，提升多文档场景下回答的可追溯性。")
    add_bullet(document, "", "针对多轮对话实现 Query Rewrite，并设计 strict / open 两种问答模式；严格模式下检索不足直接拒答，并增加答案校验环节，降低无依据回答。")
    add_bullet(document, "", "实现基于 SSE 的流式问答接口，补充语义缓存、模型懒加载和失败重试机制，优化长回答场景下的响应体验与系统稳定性。")

    p2 = document.add_paragraph()
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after = Pt(2)
    p2_run1 = p2.add_run("基于 Qt / Linux 的短视频播放器")
    set_run_font(p2_run1, size=11.5, bold=True, color=RGBColor(35, 52, 105))
    p2_run2 = p2.add_run("   Qt / Linux / FFmpeg / SDL2 / OpenGL / Nginx-RTMP / MySQL")
    set_run_font(p2_run2, size=9.6, color=RGBColor(96, 96, 96))

    add_bullet(document, "", "基于 Qt / Linux 开发桌面端视频播放器，使用 FFmpeg 完成音视频解码，结合 SDL2 处理音频播放、OpenGL 实现视频渲染。")
    add_bullet(document, "", "服务端采用 epoll + 线程池 + MySQL，配合 Nginx-RTMP 支持视频上传、推流分发与拉流播放，覆盖本地播放与网络播放场景。")
    add_bullet(document, "", "围绕音画同步和播放体验做功能完善，包括暂停/进度控制、封面抽帧与 GIF 预览等。")

    add_section_heading(document, "技术能力")
    skills_table = document.add_table(rows=0, cols=2)
    skills_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    skills_table.autofit = False
    remove_table_borders(skills_table)
    widths = [Cm(2.8), Cm(14.7)]
    skills = [
        ("C/C++", "熟悉面向对象与 STL，掌握 C++11 常用特性，具备 Linux/C++ 项目开发基础。"),
        ("后端开发", "熟悉 Python 开发，掌握 FastAPI、SQLAlchemy、RESTful 接口设计与 MySQL 基础。"),
        ("AI应用", "了解 RAG、Embedding、向量检索、提示词编排与 LLM 应用开发，能完成端到端功能落地。"),
        ("计算机基础", "掌握常用数据结构与算法，熟悉 TCP/UDP、HTTP/HTTPS、Socket 编程与 Linux 常用机制。"),
        ("工程能力", "熟悉 Git、Docker 及基础部署流程，能够使用 Claude Code、Cursor 等工具辅助开发与调试。"),
        ("实践方向", "关注检索增强问答、后端工程化与 AI 应用落地，具备独立完成课程/个人项目的能力。"),
    ]
    for label, text in skills:
        target_row = skills_table.add_row()
        label_cell = target_row.cells[0]
        text_cell = target_row.cells[1]
        label_cell.width = widths[0]
        text_cell.width = widths[1]
        label_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        text_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        lp = label_cell.paragraphs[0]
        lp.paragraph_format.space_after = Pt(1)
        lr = lp.add_run(label)
        set_run_font(lr, size=10.5, bold=True, color=RGBColor(44, 44, 44))
        tp = text_cell.paragraphs[0]
        tp.paragraph_format.space_after = Pt(2)
        tp.paragraph_format.line_spacing = 1.1
        tr = tp.add_run(text)
        set_run_font(tr, size=10.2, color=RGBColor(55, 55, 55))

    add_section_heading(document, "获奖经历")
    add_bullet(document, "", "2025 年获校级二等奖学金。")
    add_bullet(document, "", "2024 年获“挑战杯”省级三等奖。")

    return document


def main():
    document = build_document()
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
