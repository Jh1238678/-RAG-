import docx
from docx.shared import Pt

def md_to_docx(md_path, docx_path):
    doc = docx.Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        elif line.startswith('##### '):
            doc.add_heading(line[6:], level=5)
        elif line.startswith('---'):
            continue
        else:
            text = line.replace('**', '')
            doc.add_paragraph(text)
            
    doc.save(docx_path)

if __name__ == '__main__':
    md_to_docx('rewritten_doc.md', '吴凯浩_机器人23-2_改写版.docx')
    print("Convert success")
