"""Word (.docx) 解析器。使用 python-docx。

保留段落层次（标题/正文）。
"""
from app.core.logger import logger
from app.parsers.base import BaseParser, PageInfo, ParsedDocument


class DocxParser(BaseParser):
    """Word 文档解析器。"""

    def parse(self, file_path: str) -> ParsedDocument:
        from docx import Document as DocxDocument  # 延迟导入
        logger.info(f"开始解析 DOCX: {file_path}")
        doc = DocxDocument(file_path)

        paragraphs: list[str] = []
        heading_count = 0

        # 读取段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if para.style and para.style.name.startswith("Heading"):
                paragraphs.append(f"\n## {text}\n")
                heading_count += 1
            else:
                paragraphs.append(text)

        # 读取表格
        for table in doc.tables:
            paragraphs.append("\n[表格]")
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip().replace("\n", " "))
                if any(row_data):
                    paragraphs.append(" | ".join(row_data))
            paragraphs.append("[表格结束]\n")

        full_text = "\n".join(paragraphs)
        # Word 没有页码概念，整体作为一页（设为 None 避免前端显示恒为“第1页”）
        pages = [PageInfo(page_num=None, text=full_text)] if full_text else []

        logger.info(f"DOCX 解析完成: {len(paragraphs)} 段, {len(full_text)} 字符, {heading_count} 个标题")
        return ParsedDocument(
            text=full_text,
            pages=pages,
            metadata={"paragraph_count": len(paragraphs), "heading_count": heading_count, "parser": "python-docx"},
        )
